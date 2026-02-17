"""Document store for tracking and storing generated documents."""

from typing import List, Dict, Optional
from datetime import datetime
import json
import os
import io
from app.config import settings

try:
    from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas
    from azure.core.exceptions import ResourceExistsError, ResourceNotFoundError
    AZURE_AVAILABLE = True
except ImportError:
    AZURE_AVAILABLE = False
    print("Azure Storage libraries not available. Using local storage only.")

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    import markdown
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. DOCX generation disabled.")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.colors import HexColor, black, white, grey
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus.frames import Frame
    from reportlab.platypus.doctemplate import PageTemplate, BaseDocTemplate
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("reportlab not available. PDF generation disabled.")


class DocumentStore:
    """Document store for tracking and storing generated documents with Supabase or Azure Blob Storage support."""

    def __init__(self, store_file: str = "generated_documents.json"):
        """Initialize document store."""
        self.store_file = store_file
        self.store_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
            store_file
        )

        # Storage backend setup
        self.use_supabase = bool(settings.supabase_url and settings.supabase_anon_key)
        self.use_azure = bool(settings.azure_storage_connection_string)
        
        self.supabase_client = None
        self.blob_service_client = None
        self.generated_container = "generated-documents"
        self.supabase_bucket = settings.supabase_storage_bucket or "Tech_standards_bucket"

        if self.use_supabase:
            try:
                from supabase import create_client
                self.supabase_client = create_client(
                    settings.supabase_url,
                    settings.supabase_anon_key
                )
                print("[OK] Document store using Supabase Storage")
            except Exception as e:
                print(f"[WARNING] Could not initialize Supabase: {e}")
                self.use_supabase = False

        if not self.use_supabase and AZURE_AVAILABLE and settings.azure_storage_connection_string:
            try:
                self.blob_service_client = BlobServiceClient.from_connection_string(
                    settings.azure_storage_connection_string
                )
                self._ensure_container_exists()
                print("[OK] Document store using Azure Blob Storage")
            except Exception as e:
                print(f"[WARNING] Failed to initialize Azure Blob Storage: {e}")
                self.blob_service_client = None
        
        if not self.use_supabase and not self.blob_service_client:
            print("[INFO] Document store using local filesystem only")
    
    def _load_documents(self) -> List[Dict]:
        """Load documents from file."""
        if not os.path.exists(self.store_path):
            return []
        
        try:
            with open(self.store_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading document store: {e}")
            return []
    
    def _save_documents(self, documents: List[Dict]):
        """Save documents to file."""
        try:
            with open(self.store_path, 'w', encoding='utf-8') as f:
                json.dump(documents, f, indent=2, default=str)
        except Exception as e:
            print(f"Error saving document store: {e}")
    
    def add_document(self, document: Dict) -> bool:
        """Add a generated document to the store."""
        try:
            documents = self._load_documents()
            
            # Check if document already exists
            doc_id = document.get("id")
            documents = [d for d in documents if d.get("id") != doc_id]
            
            # Add new document
            document["created_at"] = datetime.utcnow().isoformat()
            documents.append(document)
            
            self._save_documents(documents)
            return True
        except Exception as e:
            print(f"Error adding document to store: {e}")
            return False
    
    def get_document(self, document_id: str) -> Optional[Dict]:
        """Get a document by ID."""
        documents = self._load_documents()
        for doc in documents:
            if doc.get("id") == document_id:
                return doc
        return None
    
    def list_documents(self) -> List[Dict]:
        """List all generated documents."""
        return self._load_documents()
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document from the store."""
        try:
            documents = self._load_documents()
            documents = [d for d in documents if d.get("id") != document_id]
            self._save_documents(documents)

            # Also delete from blob storage if available
            if self.blob_service_client:
                try:
                    blob_client = self.blob_service_client.get_blob_client(
                        container=self.generated_container,
                        blob=f"{document_id}.docx"
                    )
                    blob_client.delete_blob()
                except Exception as e:
                    print(f"Warning: Could not delete blob for {document_id}: {e}")

            return True
        except Exception as e:
            print(f"Error deleting document from store: {e}")
            return False

    def _ensure_container_exists(self):
        """Ensure the generated documents container exists."""
        if not self.blob_service_client:
            return

        try:
            self.blob_service_client.create_container(self.generated_container)
        except ResourceExistsError:
            pass  # Container already exists
        except Exception as e:
            print(f"Warning: Could not create container {self.generated_container}: {e}")

    async def store_generated_document(
        self,
        document_id: str,
        content: str,
        metadata: Dict,
        format: str = "docx"
    ) -> Optional[str]:
        """
        Store a generated document in blob storage and return download URL.

        Args:
            document_id: Unique document identifier
            content: Document content (markdown text)
            metadata: Document metadata
            format: File format (docx, pdf, markdown)

        Returns:
            Download URL if successful, None otherwise
        """
        try:
            # Generate the file content
            if format.lower() == "docx" and DOCX_AVAILABLE:
                file_content = self._generate_docx(content, metadata)
                filename = f"{document_id}.docx"
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif format.lower() == "pdf" and PDF_AVAILABLE:
                file_content = self._generate_pdf(content, metadata)
                filename = f"{document_id}.pdf"
                content_type = "application/pdf"
            else:
                # Default to markdown
                file_content = content.encode('utf-8')
                filename = f"{document_id}.md"
                content_type = "text/markdown"

            # Store in Supabase Storage if available
            if self.use_supabase and self.supabase_client:
                try:
                    # Upload to Supabase Storage
                    storage_path = f"generated/{filename}"
                    self.supabase_client.storage.from_(self.supabase_bucket).upload(
                        storage_path,
                        file_content,
                        {"content-type": content_type, "upsert": "true"}
                    )
                    
                    # Get public URL
                    download_url = self.supabase_client.storage.from_(self.supabase_bucket).get_public_url(storage_path)
                    
                    # Store metadata in Supabase table
                    self.supabase_client.table("generated_documents").upsert({
                        "id": document_id,
                        "title": metadata.get("title", ""),
                        "filename": filename,
                        "storage_path": storage_path,
                        "format": format,
                        "author": metadata.get("author", ""),
                        "document_type": metadata.get("documentType", ""),
                        "metadata": metadata
                    }).execute()
                except Exception as e:
                    print(f"[WARNING] Supabase storage failed, using local: {e}")
                    download_url = None
            # Store in Azure Blob Storage if available
            elif self.blob_service_client:
                blob_client = self.blob_service_client.get_blob_client(
                    container=self.generated_container,
                    blob=filename
                )

                blob_client.upload_blob(
                    file_content,
                    blob_type="BlockBlob",
                    content_type=content_type,
                    metadata={
                        "document_id": document_id,
                        "title": metadata.get("title", ""),
                        "author": metadata.get("author", ""),
                        "document_type": metadata.get("documentType", ""),
                        "created_at": metadata.get("uploadedAt", datetime.utcnow().isoformat())
                    },
                    overwrite=True
                )

                # Generate SAS URL for download
                download_url = self._generate_download_url(filename)
            else:
                download_url = None
            
            if download_url is None:
                # Fallback to local storage
                local_path = os.path.join(
                    os.path.dirname(self.store_path),
                    "generated_docs",
                    filename
                )
                os.makedirs(os.path.dirname(local_path), exist_ok=True)

                if isinstance(file_content, bytes):
                    with open(local_path, 'wb') as f:
                        f.write(file_content)
                else:
                    with open(local_path, 'w', encoding='utf-8') as f:
                        f.write(file_content)

                download_url = f"/api/documents/download/{document_id}"

            # Store metadata locally
            doc_record = {
                "id": document_id,
                "title": metadata.get("title", ""),
                "filename": filename,
                "download_url": download_url,
                "format": format,
                "created_at": datetime.utcnow().isoformat(),
                **metadata
            }
            self.add_document(doc_record)

            return download_url

        except Exception as e:
            print(f"Error storing generated document {document_id}: {e}")
            return None

    def _generate_docx(self, content: str, metadata: Dict) -> bytes:
        """Generate a professional branded DOCX file."""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available")

        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = DocxDocument()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Get logo path
        logo_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            "static",
            "cranswick_logo.png"
        )
        
        # Header section with logo
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Logo cell
        logo_cell = header_table.cell(0, 0)
        if os.path.exists(logo_path):
            try:
                logo_para = logo_cell.paragraphs[0]
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(2))
            except Exception as e:
                print(f"Error adding logo to DOCX: {e}")
                logo_cell.text = "CRANSWICK PLC"
        else:
            logo_cell.text = "CRANSWICK PLC"
        
        # Document info cell
        info_cell = header_table.cell(0, 1)
        info_para = info_cell.paragraphs[0]
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc_ref = metadata.get('documentReference', 'TBD')
        issue_date = metadata.get('issueDate', datetime.now().strftime('%Y-%m-%d'))
        version = metadata.get('version', '1.0')
        
        info_run = info_para.add_run(f"Document Ref: {doc_ref}\n")
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(100, 100, 100)
        
        info_run = info_para.add_run(f"Issue Date: {issue_date}\n")
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(100, 100, 100)
        
        info_run = info_para.add_run(f"Version: {version}")
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()  # Spacing
        
        # Add horizontal line
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run("_" * 80)
        run.font.color.rgb = RGBColor(107, 45, 91)  # Cranswick purple
        
        # Add title
        title = metadata.get("title", "Principle Document")
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(107, 45, 91)  # Cranswick purple
        
        # Document type
        doc_type = metadata.get('documentType', 'Principle')
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        type_run = type_para.add_run(f"DOCUMENT TYPE: {doc_type.upper()}")
        type_run.font.size = Pt(10)
        type_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Author
        author = metadata.get('author', 'Unknown')
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"Author: {author}")
        author_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # Add another horizontal line
        para = doc.add_paragraph()
        run = para.add_run("_" * 80)
        run.font.color.rgb = RGBColor(200, 200, 200)
        
        doc.add_paragraph()  # Spacing

        # Convert markdown content to docx
        lines = content.split('\n')
        current_list_style = None
        
        for line in lines:
            original_line = line
            line = line.strip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            # Remove markdown bold markers
            line = line.replace('**', '')
            
            if line.startswith('# '):
                heading = doc.add_heading(line[2:].upper(), 1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(107, 45, 91)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], 2)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(45, 45, 45)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], 3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                doc.add_paragraph(line, style='List Number')
            else:
                para = doc.add_paragraph(line)
                para.paragraph_format.space_after = Pt(6)

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.paragraph_format.space_before = Pt(20)
        run = footer_para.add_run("_" * 80)
        run.font.color.rgb = RGBColor(107, 45, 91)
        
        footer_text = doc.add_paragraph()
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_text.add_run("CONFIDENTIAL - This document is the property of Cranswick PLC\n")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run = footer_text.add_run(f"Generated by DocumentIQ | {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_pdf(self, content: str, metadata: Dict) -> bytes:
        """Generate a professional branded PDF file."""
        if not PDF_AVAILABLE:
            raise Exception("reportlab not available")

        buffer = io.BytesIO()
        
        # Custom colors - Cranswick brand
        CRANSWICK_PURPLE = HexColor('#6B2D5B')  # Deep purple
        CRANSWICK_DARK = HexColor('#2D2D2D')
        CRANSWICK_LIGHT = HexColor('#F5F5F5')
        CRANSWICK_ACCENT = HexColor('#8B4B7B')
        
        # Create custom styles
        styles = getSampleStyleSheet()
        
        # Helper to safely add style (avoid duplicate error)
        def add_style_safe(name, **kwargs):
            if name not in styles.byName:
                styles.add(ParagraphStyle(name, **kwargs))
        
        # Document Title Style
        add_style_safe(
            'DocTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=CRANSWICK_PURPLE,
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Section Heading Style
        add_style_safe(
            'SectionHeading',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=CRANSWICK_PURPLE,
            spaceBefore=16,
            spaceAfter=8,
            fontName='Helvetica-Bold',
            borderColor=CRANSWICK_PURPLE,
            borderWidth=0,
            borderPadding=0,
        )
        
        # Subsection Style
        add_style_safe(
            'SubSection',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=CRANSWICK_DARK,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold'
        )
        
        # Body Text Style - use existing BodyText but modify it
        if 'BodyText' in styles.byName:
            styles['BodyText'].fontSize = 10
            styles['BodyText'].textColor = CRANSWICK_DARK
            styles['BodyText'].alignment = TA_JUSTIFY
        else:
            add_style_safe(
                'BodyText',
                parent=styles['Normal'],
                fontSize=10,
                textColor=CRANSWICK_DARK,
                spaceBefore=4,
                spaceAfter=4,
                alignment=TA_JUSTIFY,
                fontName='Helvetica',
                leading=14
            )
        
        # Bullet Point Style
        add_style_safe(
            'BulletPoint',
            parent=styles['Normal'],
            fontSize=10,
            textColor=CRANSWICK_DARK,
            leftIndent=20,
            spaceBefore=2,
            spaceAfter=2,
            fontName='Helvetica',
            bulletIndent=10
        )
        
        # Metadata Style
        add_style_safe(
            'Metadata',
            parent=styles['Normal'],
            fontSize=9,
            textColor=grey,
            alignment=TA_RIGHT
        )
        
        # Header Info Style
        add_style_safe(
            'HeaderInfo',
            parent=styles['Normal'],
            fontSize=10,
            textColor=CRANSWICK_DARK,
            fontName='Helvetica-Bold'
        )
        
        story = []
        
        # Get logo path
        logo_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            "static",
            "cranswick_logo.png"
        )
        
        # Header with Logo and Document Info
        header_data = []
        
        # Try to add logo
        if os.path.exists(logo_path):
            try:
                logo = Image(logo_path, width=2*inch, height=0.6*inch)
                logo.hAlign = 'LEFT'
            except Exception as e:
                print(f"Error loading logo: {e}")
                logo = Paragraph("<b>CRANSWICK PLC</b>", styles['HeaderInfo'])
        else:
            logo = Paragraph("<b>CRANSWICK PLC</b>", styles['HeaderInfo'])
        
        # Document reference info
        doc_ref = metadata.get('documentReference', 'TBD')
        issue_date = metadata.get('issueDate', datetime.now().strftime('%Y-%m-%d'))
        version = metadata.get('version', '1.0')
        
        doc_info = f"""
        <b>Document Ref:</b> {doc_ref}<br/>
        <b>Issue Date:</b> {issue_date}<br/>
        <b>Version:</b> {version}
        """
        
        # Create header table
        header_table = Table(
            [[logo, Paragraph(doc_info, styles['Metadata'])]],
            colWidths=[3.5*inch, 3.5*inch]
        )
        header_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 10))
        
        # Purple divider line
        story.append(HRFlowable(
            width="100%",
            thickness=3,
            color=CRANSWICK_PURPLE,
            spaceBefore=5,
            spaceAfter=15
        ))
        
        # Document Title
        title = metadata.get("title", "Principle Document")
        story.append(Paragraph(title, styles['DocTitle']))
        
        # Document Type Badge
        doc_type = metadata.get('documentType', 'Principle')
        badge_text = f"<b>DOCUMENT TYPE:</b> {doc_type.upper()}"
        story.append(Paragraph(badge_text, styles['Metadata']))
        story.append(Spacer(1, 5))
        
        # Author and date info
        author = metadata.get('author', 'Unknown')
        info_text = f"<b>Author:</b> {author}"
        story.append(Paragraph(info_text, styles['BodyText']))
        story.append(Spacer(1, 15))
        
        # Thin divider
        story.append(HRFlowable(
            width="100%",
            thickness=1,
            color=CRANSWICK_LIGHT,
            spaceBefore=5,
            spaceAfter=15
        ))
        
        # Process content
        lines = content.split('\n')
        current_paragraph = ""
        in_table = False
        table_data = []
        
        for line in lines:
            original_line = line
            line = line.strip()
            
            if not line:
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, styles['BodyText']))
                    current_paragraph = ""
                continue
            
            # Handle markdown formatting
            # Remove ** for bold (already handled by styles)
            line = line.replace('**', '')
            
            if line.startswith('# '):
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, styles['BodyText']))
                    current_paragraph = ""
                story.append(Spacer(1, 10))
                story.append(HRFlowable(
                    width="100%",
                    thickness=2,
                    color=CRANSWICK_PURPLE,
                    spaceBefore=0,
                    spaceAfter=5
                ))
                story.append(Paragraph(line[2:].upper(), styles['SectionHeading']))
            elif line.startswith('## '):
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, styles['BodyText']))
                    current_paragraph = ""
                story.append(Paragraph(line[3:], styles['SubSection']))
            elif line.startswith('### '):
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, styles['BodyText']))
                    current_paragraph = ""
                story.append(Paragraph(f"<b>{line[4:]}</b>", styles['BodyText']))
            elif line.startswith('- ') or line.startswith('* '):
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, styles['BodyText']))
                    current_paragraph = ""
                bullet_text = f"<bullet>&bull;</bullet> {line[2:]}"
                story.append(Paragraph(bullet_text, styles['BulletPoint']))
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, styles['BodyText']))
                    current_paragraph = ""
                story.append(Paragraph(f"<b>{line}</b>", styles['BodyText']))
            elif line.startswith('|') and line.endswith('|'):
                # Table row
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, styles['BodyText']))
                    current_paragraph = ""
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if not all(c.startswith('-') for c in cells):  # Skip separator rows
                    table_data.append(cells)
            else:
                # Check if we have pending table data
                if table_data and not line.startswith('|'):
                    # Render the table
                    if len(table_data) > 0:
                        t = Table(table_data, colWidths=[2.5*inch] * len(table_data[0]) if table_data else None)
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), CRANSWICK_PURPLE),
                            ('TEXTCOLOR', (0, 0), (-1, 0), white),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                            ('TOPPADDING', (0, 0), (-1, -1), 8),
                            ('BACKGROUND', (0, 1), (-1, -1), CRANSWICK_LIGHT),
                            ('GRID', (0, 0), (-1, -1), 0.5, grey),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ]))
                        story.append(t)
                        story.append(Spacer(1, 10))
                    table_data = []
                
                current_paragraph += line + " "
        
        # Handle any remaining content
        if current_paragraph:
            story.append(Paragraph(current_paragraph, styles['BodyText']))
        
        if table_data:
            if len(table_data) > 0:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), CRANSWICK_PURPLE),
                    ('TEXTCOLOR', (0, 0), (-1, 0), white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), CRANSWICK_LIGHT),
                    ('GRID', (0, 0), (-1, -1), 0.5, grey),
                ]))
                story.append(t)
        
        # Footer section
        story.append(Spacer(1, 30))
        story.append(HRFlowable(
            width="100%",
            thickness=2,
            color=CRANSWICK_PURPLE,
            spaceBefore=10,
            spaceAfter=10
        ))
        
        footer_text = f"""
        <b>CONFIDENTIAL</b> - This document is the property of Cranswick PLC<br/>
        Generated by DocumentIQ | {datetime.now().strftime('%Y-%m-%d %H:%M')}
        """
        story.append(Paragraph(footer_text, styles['Metadata']))
        
        # Build PDF with page numbers
        def add_page_number(canvas, doc):
            page_num = canvas.getPageNumber()
            text = f"Page {page_num}"
            canvas.saveState()
            canvas.setFont('Helvetica', 8)
            canvas.setFillColor(grey)
            canvas.drawRightString(7.5*inch, 0.5*inch, text)
            canvas.restoreState()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
        
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_download_url(self, filename: str) -> str:
        """Generate a SAS URL for downloading the blob."""
        if not self.blob_service_client:
            return f"/api/generate/download/{filename}"

        try:
            from datetime import datetime, timedelta
            from azure.storage.blob import BlobSasPermissions, generate_blob_sas

            # Generate SAS token valid for 24 hours
            sas_token = generate_blob_sas(
                account_name=self.blob_service_client.account_name,
                container_name=self.generated_container,
                blob_name=filename,
                account_key=self.blob_service_client.credential.account_key,
                permission=BlobSasPermissions(read=True),
                expiry=datetime.utcnow() + timedelta(hours=24)
            )

            return f"{self.blob_service_client.primary_endpoint}{self.generated_container}/{filename}?{sas_token}"

        except Exception as e:
            print(f"Error generating SAS URL: {e}")
            return f"/api/generate/download/{filename}"

    async def get_download_url(self, document_id: str) -> Optional[str]:
        """Get download URL for a generated document."""
        # First check local store
        doc = self.get_document(document_id)
        if doc and doc.get("download_url"):
            return doc["download_url"]

        # If not found locally and blob storage available, try to generate URL
        if self.blob_service_client:
            try:
                # Check if blob exists
                blob_client = self.blob_service_client.get_blob_client(
                    container=self.generated_container,
                    blob=f"{document_id}.docx"
                )

                if blob_client.exists():
                    return self._generate_download_url(f"{document_id}.docx")

                # Try other formats
                for ext in ['.pdf', '.md']:
                    blob_client = self.blob_service_client.get_blob_client(
                        container=self.generated_container,
                        blob=f"{document_id}{ext}"
                    )
                    if blob_client.exists():
                        return self._generate_download_url(f"{document_id}{ext}")

            except Exception as e:
                print(f"Error checking blob existence: {e}")

        return None


# Global instance
document_store = DocumentStore()
